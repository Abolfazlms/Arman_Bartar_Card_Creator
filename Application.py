import os 
import pandas as pd
from tkinter import *
from tkinter import filedialog
from docx import Document 
from docx2pdf import convert
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, RGBColor, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT

file_path = None  # The Path of File
term_number = None

def select_file():
    global file_path
    # Open The File Selection Window
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    if file_path: 
        entry.delete(0, END) # Remove The Text On The Box
        entry.insert(0, file_path) # Display The FilePath On The Box
def get_status(score):
    """Determine the status based on the score."""
    if score >= 17 and score <= 20:
        return "عالی"
    elif score < 17 and score >= 14:
        return "بسیار خوب"
    elif  score < 14 and score >= 12:
        return "قابل قبول"
    else:
        return "مردود"

def apply_font_and_color_to_paragraph(paragraph, font_name='B Nazanin', font_size=12, font_color=RGBColor(0, 0, 0), font_bold=False):

    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = font_color
        run.font.bold = font_bold

def process_file():
    global file_path, term_number
    if not file_path or not term_number:
        # Check If The Excel File Is Selected.
        result_label.config(text="Select a file and Enter Term Number Please!", fg="red")
        return
    try:        
        # Reading Excel File
        df = pd.read_excel(file_path, engine='openpyxl')
        # print(df)

        # Desktop File Save Directory
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        folder_name = "Report Card" 
        folder_path = os.path.join(desktop_path, folder_name)

        # Create Folder If Didn't Exist.
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        # Find Name of The Teachers From ExcelFile
        teacher_columns = df.columns[-2:]  # Two Last Column
        score_columns = df.columns[3:5]   # Grades Column
        courses = list(score_columns)
        teachers = list(teacher_columns)

        # Create a Word File For Each Record (Row)
        for _, row in df.iterrows():
            student_id = row['ردیف']
            name = row['نام و نام خانوادگی']
            value_number = row['شماره ارزش آفرینی']
            scores = {course: row[course] for course in courses}
            teacher_names = {course: row[teacher] for course, teacher in zip(courses, teachers)}
            
            # Create A Document For Each Person
            doc = Document()

            # Landscape Orientation
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE  
            # Paper Size
            section.page_width = Inches(8.26) 
            section.page_height = Inches(5.82)

            # Set The Margins Of Document
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.2)
            section.bottom_margin = Inches(0.25)

            # Change The Direction and Alignment Of Document To Right
            for paragraph in doc.paragraphs:
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT 
                paragraph.paragraph_format.bidi = True 

            # Read The Logo Path
            image_path = os.path.join(os.getcwd(), "logo.png")
            # Add Logo to Document
            doc.add_picture(image_path, width=Inches(0.7))

            # Add Text On Document
            text = doc.add_paragraph("بسمه تعالی",)
            text.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_font_and_color_to_paragraph(text, font_name='B Nazanin', font_size=12, font_bold=True)

            text = doc.add_paragraph(f'" کارنامه ترم {term_number} آرمان برتر "',)
            text.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_font_and_color_to_paragraph(text, font_name='B Nazanin', font_size=12, font_bold=True)
            
            text = doc.add_paragraph(f"نام و نام خانوادگی: {name}")
            text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            apply_font_and_color_to_paragraph(text, font_name='B Nazanin', font_size=12)

            text = doc.add_paragraph(f"شماره ارزش آفرینی: {value_number}")
            text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            apply_font_and_color_to_paragraph(text, font_name='B Nazanin', font_size=12)


            # Add Table To Document
            table = doc.add_table(rows=1, cols=4) 
            table.style = 'Medium Shading 1 Accent 1'
            def set_column_width(table, col_idx, width_in_inch):
                for row in table.rows:
                    cell = row.cells[col_idx]
                    cell.width = Inches(width_in_inch)  # Change The Width of Column To Inch

            table.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Format The Headers Of Table
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "وضعیت"
            hdr_cells[1].text = "نمره"
            hdr_cells[2].text = "نام استاد"
            hdr_cells[3].text = "نام درس"

            for cell in hdr_cells:
                # Format The Text In Table Cells
                run = cell.paragraphs[0].runs[0]
                run.font.name = 'B Nazanin'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            for cell in hdr_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
            # Add The Rows on Table
            for course in courses:
                score = scores[course]
                row = table.add_row().cells
                row[0].text = get_status(score)  # وضعیت بر اساس نمره
                row[1].text = str(score)
                row[2].text = teacher_names[course]
                row[3].text = course

                # Format The Text On The Table Rows
                for cell in row:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = 'B Nazanin'
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)

            # Apply borders to all cells in the table
            for row in table.rows:
                set_column_width(table, row._index, 1.5)
                for cell in row.cells:
                    # Add borders to all sides of the cell
                    cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders %s>'  
                        r'<w:top w:val="single" w:space="0" w:size="4" w:color="000000"/>'
                        r'<w:left w:val="single" w:space="0" w:size="4" w:color="000000"/>'
                        r'<w:bottom w:val="single" w:space="0" w:size="4" w:color="000000"/>'
                        r'<w:right w:val="single" w:space="0" w:size="4" w:color="000000"/>'
                        r'</w:tcBorders>' % nsdecls('w')))
            
            # Add The Guide For Score
            text = doc.add_paragraph("", style='Normal')        
            text = doc.add_paragraph("عالی: معادل ۱۷ الی ۲۰*", style='Normal')
            text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            apply_font_and_color_to_paragraph(text, font_name='B Nazanin', font_size=10, font_bold=True)

            text = doc.add_paragraph("بسیار خوب: معادل ۱۴ الی ۱۷", style='Normal')
            text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            apply_font_and_color_to_paragraph(text, font_name='B Nazanin', font_size=10, font_bold=True)

            text = doc.add_paragraph("قابل قبول: معادل ۱۲ الی ۱۴", style='Normal')
            text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            apply_font_and_color_to_paragraph(text, font_name='B Nazanin', font_size=10, font_bold=True)

            text = doc.add_paragraph("مردود: معادل کمتر از ۱۲", style='Normal')
            text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            apply_font_and_color_to_paragraph(text, font_name='B Nazanin', font_size=10, font_bold=True)

        
            # Save The WordFile In App Directory
            # file_name = f"student_{student_id}.docx"
            # doc.save(file_name)

            # Save The File On Custom Folder on Desktop
            file_name = f"{name}.docx"
            docx_path = os.path.join(folder_path, file_name)  # Full Path of File
            doc.save(docx_path)

            # Set The Path For Pdf Format Export
            pdf_folder_path = os.path.join(desktop_path, "PDF Export")
            
            # Create Folder If Didn't Exist.
            if not os.path.exists(pdf_folder_path):
                os.makedirs(pdf_folder_path)

            # Convert and Export The .docx File to .pdf File
            convert(docx_path, pdf_folder_path)
        
        # Show Successful Processing Message.
        result_label.config(text="The operation was successful!", fg="green")
    except Exception as e:
        # Show The Exception Message
        result_label.config(text=f"Error: {str(e)}", fg="red")
        
# Create The GUI of Mini Application
root = Tk()

# Find The Desktop Page Resolution
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# Calculating The Coordinates of Start Point of Application
x_position = screen_width - root.winfo_reqwidth() - 500  # Distance From Right Edge
y_position = 50  # Distance From Top Edge

# Set The Window Position
root.geometry(f"+{x_position}+{y_position}")

# Widget Section
root.title("Arman Bartar Mini App")
# Create Entry Box to Display FilePath
entry = Entry(root, width=50)
entry.grid(row=0, column=0, padx=10, pady=10)

# Create Button For File Selection
button = Button(root, text="Select File", command=select_file)
button.grid(row=0, column=1, padx=10, pady=10)

# Create Entry Box for Term Number
term_entry = Entry(root, width=50)
term_entry.grid(row=1, column=0, padx=10, pady=10)
term_label = Label(root, text="ترم (به حروف)")
term_label.grid(row=1, column=1, padx=10, pady=10)

# Set the term_number variable when the user clicks process
def set_term_number():
    global term_number
    term_number = term_entry.get()  # Get the term number from the entry field

# Create Process Button
process_button = Button(root, text="Process", command=lambda: [set_term_number(), process_file()])
process_button.grid(row=2, column=0, padx=10, pady=10, sticky="ew")

# Exit Button
exit_button = Button(root, text="Exit", command=root.quit)
exit_button.grid(row=2, column=1, padx=10, pady=10, sticky="ew")


# Create Result Label
result_label = Label(root, text="", font=("Arial", 10))
result_label.grid(row=3, column=0, columnspan=2, pady=5)

# End of Widget Section
root.mainloop()
